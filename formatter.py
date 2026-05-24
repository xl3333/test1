""""排版模块"""
import re
from docx import Document
from docx.shared import Pt,Inches
from docx.enum.text import WD_LINE_SPACING
class DocumentFormatter:
    DEFAULT_FONT = '宋体'
    DEFAULT_FONT_SIZE = 12
    DEFAULT_LINE_SPACING = 18
    DEFAULT_INDENT = 0.5
    def __init__(self):
        self.custom_settings={
            'font_name':self.DEFAULT_FONT,
            'font_size':self.DEFAULT_FONT_SIZE,
            'line_spacing':self.DEFAULT_LINE_SPACING,
            'indent':self.DEFAULT_INDENT,
        }
    def clean_whitespace(self,content):
        content = re.sub(r'\s+','',content)
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        cleaned_lines=[]
        empty_count=0
        for line in lines:
            if line=='':
                empty_count+=1
                if empty_count<=2:
                    cleaned_lines.append(line)
                else:
                    empty_count=0
                    cleaned_lines.append(line)
        while cleaned_lines and cleaned_lines[0]=='':
            cleaned_lines.pop(0)
        while cleaned_lines and cleaned_lines[-1]=='':
            cleaned_lines.pop(-1)
        return '\n'.join(cleaned_lines)
    def save_as_word(self,content,out_path):
        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = self.custom_settings['font_name']
            style.font.size = Pt(self.custom_settings['font_size'])
            style.paragraph_format.line_spacing = Pt(self.custom_settings['line_spacing'])
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraphs = content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    paragraph = doc.add_paragraph(para_text)
                    paragraph.style = style
                elif para_text.strip() and not para_text.strip().startswith('*/'):
                    paragraph = doc.add_paragraph(para_text)
                    paragraph.style = style
            doc.save(out_path)
            return True
        except Exception as e:
            raise e(f"保存失败: {out_path}")



""""文件处理模块"""
import os
from itertools import count

from docx import Document
class FileHandler:
    type_map =['.docx','md','.txt']
    def read_file(file_path):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        file_type = os.path.splitext(file_path)[1]
        if file_type not in FileHandler.type_map:
            raise ValueError(f"文件类型不支持: {file_type}")
        try:
            if file_type in['.txt','.md']:
                with open(file_path,'r',encoding='utf-8') as f:
                    content = f.read()
                return content,'txt'
            elif file_type == '.docx':
                doc = Document(file_path)
                content =[]
                for paragraph in doc.paragraphs:
                    content.append(paragraph.text)
                    return '\n'.join(content),'docx'
        except Exception as e:
            raise e(f"文件读取失败: {file_path}")
    def save_file(content,file_path,file_type):
        try:
            if file_type == 'txt':
                with open(file_path,'w',encoding='utf-8')as f:
                    f.write(content)
            elif file_type == 'docx':
                doc = Document()
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
                doc.save(file_path)
        except Exception as e:
            raise e(f"文件保存失败: {file_path}")






